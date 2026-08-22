#!/usr/bin/env python3
"""PHASE D — audit the recorded answers in testing_results/.

WHY THIS RUNS FIRST, BEFORE ANY FIXES
    It needs no code changes, no re-ingestion, no new queries, and no API
    calls. It reads answers that were already produced and asks one question:

        did the answer cite the provisions the system actually retrieved?

    If the answer is "mostly yes", the paper's premise is wrong and we find out
    in week one instead of November. See docs/EVALUATION_PLAN.md.

WHAT THE RECORDED FILES CONTAIN
    Four .docx files. Two captured the "Citations / Sources" panel; two did
    not. So:

        runs with a panel -> full audit (grounded / ungrounded / divergence)
        runs without      -> out-of-corpus rate only

    Out-of-corpus is still meaningful without a retrieval record: if an answer
    cites the Negotiable Instruments Act and that Act was never indexed, the
    model cannot have retrieved it, whatever the retrieval was.

A PARSING NOTE THAT MATTERS FOR CORRECTNESS
    python-docx exposes doc.paragraphs and doc.tables as two separate lists,
    which loses their relative order. The "Relevant Legal Provisions" table
    belongs to the answer immediately above it. Reading the lists separately
    and concatenating attributes EVERY table to EVERY run and inflates the
    citation count roughly sevenfold. So we walk the document body in order.

IMPORTANT CAVEAT
    These runs predate commit ce1bded (the confidence system). They evaluate a
    code version that no longer exists. Treat every number as a LABELLED
    PRE-FIX BASELINE, never as an evaluation of the current system.

USAGE
    python scripts/audit_recorded_runs.py
    python scripts/audit_recorded_runs.py --json runs/phase_d_audit.json
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))
sys.stdout.reconfigure(encoding="utf-8")

from docx import Document  # noqa: E402
from docx.table import Table  # noqa: E402
from docx.text.paragraph import Paragraph  # noqa: E402

from core.verifier import CorpusIndex, audit_answer  # noqa: E402

ROOT = Path(__file__).resolve().parents[1]
RESULTS_DIR = ROOT / "testing_results"

# What is actually in the Qdrant index — four documents, four Acts.
# Anything cited outside this set could not possibly have been retrieved.
CORPUS_ACTS = {"IPC", "BNS", "CRPC", "CPA"}

QUESTION_MARK = "\U0001F9D1"  # the 🧑‍⚖️ emoji prefixing each recorded question
CITATION_RE = re.compile(r"^\s*\[\d+\]\s+\S+\.pdf,\s*p\.", re.IGNORECASE)
DASHBOARD_MARKERS = ("Paralegal Dashboard", "Citations / Sources", "Raw JSON")

NL = "\n"


def ordered_blocks(doc: Document) -> list[str]:
    """Walk the document body in order, flattening paragraphs and tables alike."""
    out: list[str] = []
    for child in doc.element.body.iterchildren():
        tag = child.tag.split("}")[-1]
        if tag == "p":
            out.append(Paragraph(child, doc).text)
        elif tag == "tbl":
            table = Table(child, doc)
            rows = [" | ".join(c.text.strip() for c in r.cells) for r in table.rows]
            out.append(NL.join(rows))
    return out


def parse_doc(path: Path) -> list[dict]:
    """Split one .docx into runs, preserving table position."""
    doc = Document(path)
    blocks = ordered_blocks(doc)
    marks = [i for i, t in enumerate(blocks) if t.strip().startswith(QUESTION_MARK)]

    if marks:
        runs = []
        for k, start in enumerate(marks):
            end = marks[k + 1] if k + 1 < len(marks) else len(blocks)
            block = blocks[start:end]
            question = block[0].split("Question:", 1)[-1].strip()

            # The answer ends where the dashboard / citations panel begins.
            # Any table inside that span belongs to THIS answer.
            stop = next(
                (i for i, t in enumerate(block) if any(m in t for m in DASHBOARD_MARKERS)),
                len(block),
            )
            answer = NL.join(block[1:stop])
            panel = NL.join(t for t in block if CITATION_RE.match(t.strip()))
            runs.append(
                {
                    "source": path.name,
                    "question": question,
                    "answer": answer,
                    "panel": panel,
                    "has_retrieval_record": bool(panel.strip()),
                }
            )
        return runs

    # No question markers: treat the whole document as one block. Out-of-corpus
    # citations remain measurable without a retrieval record.
    panel = NL.join(t for t in blocks if CITATION_RE.match(t.strip()))
    return [
        {
            "source": path.name,
            "question": f"({path.stem}: questions not captured in this file)",
            "answer": NL.join(blocks),
            "panel": panel,
            "has_retrieval_record": bool(panel.strip()),
        }
    ]


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("--json", type=Path, help="also write full results here")
    args = ap.parse_args()

    corpus = CorpusIndex(acts=CORPUS_ACTS, sections=[])
    files = sorted(RESULTS_DIR.glob("*.docx"))
    if not files:
        print(f"No .docx files in {RESULTS_DIR}")
        return

    rows, full = [], []
    for f in files:
        for run in parse_doc(f):
            res = audit_answer(
                answer_text=run["answer"],
                retrieved_text=run["panel"],
                corpus=corpus,
                panel_text=run["panel"],
            )
            rows.append((run, res))
            full.append(
                {
                    **{k: v for k, v in run.items() if k != "answer"},
                    "cited": [str(p) for p in res.cited],
                    "grounded": [str(p) for p in res.grounded],
                    "ungrounded": [str(p) for p in res.ungrounded],
                    "out_of_corpus": [str(p) for p in res.out_of_corpus],
                    "vintage_errors": res.vintage_errors,
                    **res.summary(),
                }
            )

    print("=" * 102)
    print("PHASE D — CITATION AUDIT OF RECORDED RUNS")
    print("=" * 102)
    print()
    print("CAVEAT: these runs predate the confidence system (commit ce1bded).")
    print("        Treat every number as a labelled PRE-FIX BASELINE.")
    print()
    print(f"{'QUERY':<46}{'CITED':>6}{'GRND':>6}{'UNGR':>6}{'OOC':>5}{'JACC':>7}  RETRIEVAL?")
    print("-" * 102)
    for run, res in rows:
        s = res.summary()
        print(
            f"{run['question'][:44]:<46}{s['n_cited']:>6}{s['n_grounded']:>6}"
            f"{s['n_ungrounded']:>6}{s['n_out_of_corpus']:>5}"
            f"{s['panel_prose_jaccard']:>7.2f}"
            f"  {'yes' if run['has_retrieval_record'] else 'NO RECORD'}"
        )

    with_ret = [(r, a) for r, a in rows if r["has_retrieval_record"]]
    total_cited = sum(a.n_cited for _, a in rows)
    total_ooc = sum(len(a.out_of_corpus) for _, a in rows)

    print()
    print("=" * 102)
    print("AGGREGATE")
    print("=" * 102)
    print(f"  runs audited                      {len(rows)}")
    print(f"  runs with a retrieval record      {len(with_ret)}")
    print(f"  total provisions cited            {total_cited}")

    if total_cited:
        print()
        print(f"  OUT-OF-CORPUS RATE  (all {len(rows)} runs)")
        print(
            f"    cited an Act never indexed      {total_ooc}/{total_cited}"
            f"  = {100 * total_ooc / total_cited:.1f}%"
        )

    if with_ret:
        c = sum(a.n_cited for _, a in with_ret)
        g = sum(len(a.grounded) for _, a in with_ret)
        u = sum(len(a.ungrounded) for _, a in with_ret)
        j = sum(a.panel_prose_jaccard for _, a in with_ret) / len(with_ret)
        v = sum(len(a.vintage_errors) for _, a in with_ret)
        none_grounded = sum(1 for _, a in with_ret if a.n_cited and not a.grounded)
        print()
        print(f"  GROUNDING  ({len(with_ret)} runs with a retrieval record)")
        print(f"    provisions cited                {c}")
        if c:
            print(f"    grounded (found in retrieval)   {g}   = {100 * g / c:.1f}%")
            print(f"    UNGROUNDED                      {u}   = {100 * u / c:.1f}%")
        print(f"    mean panel-prose Jaccard        {j:.3f}   (1.0 = perfect agreement)")
        print(f"    vintage errors (wrong scheme)   {v}")
        print(f"    answers citing NOTHING that was retrieved:  {none_grounded}/{len(with_ret)}")

    if args.json:
        args.json.parent.mkdir(parents=True, exist_ok=True)
        args.json.write_text(
            json.dumps(full, indent=2, ensure_ascii=False), encoding="utf-8"
        )
        print()
        print(f"  full per-run results -> {args.json}")


if __name__ == "__main__":
    main()
