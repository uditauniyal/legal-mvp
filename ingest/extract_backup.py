import fitz, pdfplumber, pytesseract, re, os, io, shutil
from PIL import Image
from docx import Document
from pathlib import Path
from core.config import LANGS_OCR
from langdetect import detect
import uuid

# In-memory byte-based extraction functions (Windows-safe)
def _tesseract_ok():
    return shutil.which("tesseract") is not None

def _pm_to_pil(pix):
    mode = "RGBA" if pix.alpha else "RGB"
    return Image.frombytes(mode, [pix.width, pix.height], pix.samples)

def extract_text_pdf_bytes(data: bytes) -> list[tuple[int, str]]:
    doc = fitz.open(stream=data, filetype="pdf")
    out = []
    for i in range(doc.page_count):
        page = doc.load_page(i)
        t = page.get_text("text") or ""
        if len(t.strip()) < 20 and _tesseract_ok():
            pix = page.get_pixmap(dpi=300)
            img = _pm_to_pil(pix)
            t = pytesseract.image_to_string(img, lang=LANGS_OCR)
        out.append((i+1, t))
    return out

def extract_text_docx_bytes(data: bytes) -> list[tuple[int, str]]:
    doc = Document(io.BytesIO(data))
    text = "\n".join(p.text for p in doc.paragraphs)
    return [(1, text)]

def extract_text_txt_bytes(data: bytes) -> list[tuple[int, str]]:
    return [(1, data.decode("utf-8", errors="ignore"))]

# Original functions for reference
def extract_text_pdf(path: str) -> list[tuple[int, str]]:
    pages = []
    with pdfplumber.open(path) as pdf:
        for i, p in enumerate(pdf.pages, start=1):
            t = p.extract_text() or ""
            if len(t.strip()) < 20:
                try:
                    pix = fitz.open(path)[i-1].get_pixmap(dpi=300)
                    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                    t = pytesseract.image_to_string(img, lang=LANGS_OCR)
                except Exception as e:
                    print(f"OCR failed for page {i}: {e}")
                    t = ""
            pages.append((i, t))
    return pages

def extract_text_docx(path: str) -> list[tuple[int, str]]:
    doc = Document(path)
    text = "\n".join(p.text for p in doc.paragraphs)
    return [(1, text)]

def extract_text_txt(path: str) -> list[tuple[int, str]]:
    return [(1, Path(path).read_text(encoding="utf-8", errors="ignore"))]